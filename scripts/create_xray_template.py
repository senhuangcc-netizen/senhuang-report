"""
生成 X光報告 xray_base.docx 模板（docxtemplater 格式）
執行：python3 scripts/create_xray_template.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent.parent / "templates" / "xray_base.docx"
FONT = "PMingLiU"


def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(Cm(height_cm).emu / 914400 * 1440)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)


def add_run(para, text, bold=False, size=None, color=None, underline=False):
    r = para.add_run(text)
    r.font.name = FONT
    r.font.size = size or Pt(11)
    r.font.bold = bold
    r.font.underline = underline
    if color:
        r.font.color.rgb = color
    try:
        r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    except Exception:
        pass
    return r


doc = Document()
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(3.17)
section.top_margin  = section.bottom_margin = Cm(2.54)


def dp(text='', bold=False, size=None, align=WD_ALIGN_PARAGRAPH.LEFT,
       underline=False, space_before=Pt(0), space_after=Pt(3)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    if text:
        add_run(p, text, bold=bold, size=size, underline=underline)
    return p


# 標題
dp('東方森煌古物鑑定所 X-RAY 顯影圖', bold=True, size=Pt(14),
   align=WD_ALIGN_PARAGRAPH.CENTER, underline=True, space_after=Pt(2))
dp('Asia SenHuang Authentication X-ray Radiograph',
   align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(11), space_after=Pt(4))

# 編號 + 日期
p_info = doc.add_paragraph()
p_info.paragraph_format.space_before = Pt(0)
p_info.paragraph_format.space_after  = Pt(4)
add_run(p_info, '編號NO：', size=Pt(10))
add_run(p_info, '{xray_code}', size=Pt(10))
add_run(p_info, '                  ', size=Pt(10))
add_run(p_info, '日期S Date：', size=Pt(10))
add_run(p_info, '{date}', size=Pt(10))

# Table 1: Item + Angle checkboxes
COL_W = Cm(14.66)
tbl1 = doc.add_table(rows=2, cols=1)
tbl1.style = 'Table Grid'
for row in tbl1.rows:
    row.cells[0].width = COL_W

# Item row
p0 = tbl1.rows[0].cells[0].paragraphs[0]
p0.paragraph_format.space_before = Pt(2)
p0.paragraph_format.space_after  = Pt(2)
add_run(p0, 'Item: ', size=Pt(13))
add_run(p0, '{item_line}', size=Pt(13))

# Angle row
p1 = tbl1.rows[1].cells[0].paragraphs[0]
p1.paragraph_format.space_before = Pt(2)
p1.paragraph_format.space_after  = Pt(2)
add_run(p1, 'Angle: ', size=Pt(13))
add_run(p1, '{angle_line}', size=Pt(13))

# 顯影圖標題
dp('顯影圖 Item Radiograph：', size=Pt(11), space_before=Pt(4), space_after=Pt(2))

# Table 2: photo + note
tbl2 = doc.add_table(rows=2, cols=1)
tbl2.style = 'Table Grid'
for row in tbl2.rows:
    row.cells[0].width = COL_W

# Photo row（高度留給 X 光圖）
set_row_height(tbl2.rows[0], 17.0)
cell_photo = tbl2.rows[0].cells[0]
p_photo = cell_photo.paragraphs[0]
p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_photo.paragraph_format.space_before = Pt(0)
p_photo.paragraph_format.space_after  = Pt(0)
add_run(p_photo, '{%photo_xray}')
cell_photo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Note row
p_note = tbl2.rows[1].cells[0].paragraphs[0]
p_note.paragraph_format.space_before = Pt(2)
p_note.paragraph_format.space_after  = Pt(2)
add_run(p_note, '備註Note：', size=Pt(11))
add_run(p_note, '{note}', size=Pt(11))

# Footer
p_f = dp(space_before=Pt(6), space_after=Pt(0))
p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_f,
        'TEL: +8862-82602664  ｜  Web: www.senhuang.org  ｜  Email: info@senhuang.org',
        size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))

Path(OUT).parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"✓ X光模板已生成：{OUT}")
