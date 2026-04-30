"""
生成 base.docx 模板（純文字版，docxtemplater 格式）
執行：python3 scripts/create_template.py
欄寬對應 Google Doc 正確版本：50.5% / 25.2% / 24.3%
"""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent.parent / "templates" / "base.docx"
FONT      = "PMingLiU"
FONT_KAITI = "BiauKai"
PT11 = Pt(11)
GRAY = RGBColor(0xaa, 0xaa, 0xaa)

# 欄寬對應 Google Doc 正確版本（單位 pt：233.2 / 116.7 / 112.2）
COL1 = Cm(8.23)   # 50.5%  主體照 / 材質 / 說明
COL2 = Cm(4.12)   # 25.2%
COL3 = Cm(3.96)   # 24.3%  重量 / 備註 (獨立單格)
ROW_PHOTO1 = Cm(7.00)
ROW_PHOTO2 = Cm(6.70)


def set_row_height(row, height):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height.emu / 914400 * 1440)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)


def set_cell_margins_zero(cell):
    """照片格：移除 Word 預設 cell margin，讓圖片可以滿版填格"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ('top', 'start', 'bottom', 'end'):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), '0')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def fix_table_col_widths(table, widths_cm):
    """
    直接覆寫 w:tblGrid（Word 以此為準）及每個 cell 的 w:tcW。
    python-docx cell.width= 只改 tcW、不改 tblGrid，所以 Word 仍用等分預設值。
    """
    widths_tw = [int(w / 2.54 * 1440) for w in widths_cm]
    tbl = table._tbl

    old = tbl.find(qn('w:tblGrid'))
    if old is not None:
        tbl.remove(old)
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths_tw:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)

    for row in table.rows:
        col = 0
        for tc in row._tr.findall(qn('w:tc')):
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            gs_el = tcPr.find(qn('w:gridSpan'))
            span = int(gs_el.get(qn('w:val'), 1)) if gs_el is not None else 1
            cell_w = sum(widths_tw[col:col + span])
            col += span
            for ex in tcPr.findall(qn('w:tcW')):
                tcPr.remove(ex)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(cell_w))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)


def run(para, text, bold=False, size=None, color=None, italic=False, font=None):
    r = para.add_run(text)
    f = font or FONT
    r.font.name = f
    r.font.size = size or PT11
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    r._element.rPr.rFonts.set(qn('w:eastAsia'), f)
    return r


def cell_label_value(cell, zh_label, en_label, value_tag='', bold_label=False, value_font=None):
    """標題置左（中文黑色 + 英文灰色），內文置中獨立段落"""
    cell.paragraphs[0].clear()
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after = Pt(1)
    run(p1, zh_label, bold=bold_label)
    run(p1, en_label, bold=bold_label, color=GRAY)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after = Pt(2)
    if value_tag:
        run(p2, value_tag, font=value_font)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def cell_text(cell, label, value_tag='', bold_label=False, center=False, value_font=None):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if label:
        run(p, label, bold=bold_label)
    if value_tag:
        run(p, value_tag, font=value_font)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def cell_photo(cell, tag):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run(p, tag)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_margins_zero(cell)


doc = Document()
section = doc.sections[0]
section.page_width   = Cm(21)
section.page_height  = Cm(29.7)
section.left_margin  = section.right_margin  = Cm(2.54)
section.top_margin   = section.bottom_margin = Cm(2.54)


def dp(text='', bold=False, size=None, align=WD_ALIGN_PARAGRAPH.LEFT,
       underline=False, space_after=Pt(2)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = space_after
    if text:
        r = run(p, text, bold=bold, size=size)
        r.font.underline = underline
    return p


def dp_label(parts, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(2)):
    """
    parts: list of (text, gray) tuples.
    冒號前英文 gray=True，其餘 gray=False。
    """
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = space_after
    for text, gray in parts:
        run(p, text, color=GRAY if gray else None)
    return p


def add_right_tab(para, pos_cm):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(pos_cm / 2.54 * 1440)))
    tabs.append(tab)
    pPr.append(tabs)


# ── 標題 ──
dp('東方森煌古物鑑定所檢驗報告', bold=True, size=Pt(13),
   align=WD_ALIGN_PARAGRAPH.CENTER, underline=True)
dp('Asia SenHuang Authentication Analysis Report', bold=True, size=Pt(11),
   align=WD_ALIGN_PARAGRAPH.CENTER)

dp_label([('送驗編號 ', False), ('NO', True), ('：{item_code}', False)])

# 送檢日期（左）與報告日期（右）同一行，以右對齊 tab 分隔
p_dates = doc.add_paragraph()
p_dates.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_dates.paragraph_format.space_before = Pt(1)
p_dates.paragraph_format.space_after  = Pt(2)
add_right_tab(p_dates, 15.92)   # 頁面內容寬度 21-2*2.54=15.92cm
run(p_dates, '送檢日期 ')
run(p_dates, 'S Date', color=GRAY)
run(p_dates, '：{submission_date}')
run(p_dates, '\t')
run(p_dates, '報告日期 ')
run(p_dates, 'R Date', color=GRAY)
run(p_dates, '：{report_date}')

dp_label([('顧客推估年代/形制 ', False), ('Presumed by customers', True), ('：{presumed}', False)])
dp_label([('送檢相關圖片 ', False), ('Item Pix', True), ('：', False)])

# ── 主表格 ──
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
for row in table.rows:
    for i, w in enumerate([COL1, COL2, COL3]):
        row.cells[i].width = w

# Row 0：主體照 | XRF（col2+3 合併）
set_row_height(table.rows[0], ROW_PHOTO1)
table.rows[0].cells[1].merge(table.rows[0].cells[2])
cell_photo(table.rows[0].cells[0], '{%photo_front}')
cell_photo(table.rows[0].cells[1], '{%photo_xrf}')

# Row 1：顯微照1 | 顯微照2（col2+3 合併）
set_row_height(table.rows[1], ROW_PHOTO2)
table.rows[1].cells[1].merge(table.rows[1].cells[2])
cell_photo(table.rows[1].cells[0], '{%photo_micro1}')
cell_photo(table.rows[1].cells[1], '{%photo_micro2}')

# Row 2：尺寸（col1+2 合併）| 重量（標題置左，內文置中）
table.rows[2].cells[0].merge(table.rows[2].cells[1])
cell_label_value(table.rows[2].cells[0], '尺寸', ' Size：', '{size}')
cell_label_value(table.rows[2].cells[2], '重量', ' Weight：gram', '{weight}')

# Row 3：材質 | 形制（col2+3 合併，標題置左，內文置中）
table.rows[3].cells[1].merge(table.rows[3].cells[2])
cell_label_value(table.rows[3].cells[0], '材質', ' Material：', '{material}')
cell_label_value(table.rows[3].cells[1], '形制', ' Category：', '{category}')

# Row 4：鑑定說明（全合併），說明文字用標楷體
table.rows[4].cells[0].merge(table.rows[4].cells[1])
table.rows[4].cells[0].merge(table.rows[4].cells[2])
cell_text(table.rows[4].cells[0], '鑑定說明 Description：\n', '{description}',
          value_font=FONT_KAITI)

# Row 5：鑑定結果（col1+2 合併）| 備註
table.rows[5].cells[0].merge(table.rows[5].cells[1])
cell_text(table.rows[5].cells[0], '鑑定結果 Result：\n', '{result}', bold_label=True)
cell_text(table.rows[5].cells[2], '備註 Note：\n', '{note}')

# ── 頁腳 ──
p_f = dp(space_after=Pt(0))
p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p_f._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top = OxmlElement('w:top')
top.set(qn('w:val'), 'single')
top.set(qn('w:sz'), '4')
top.set(qn('w:color'), 'DDDDDD')
pBdr.append(top)
pPr.append(pBdr)
p_f.paragraph_format.space_before = Pt(6)
run(p_f, 'TEL: +8862-82602664  ｜  Web: www.senhuang.org  ｜  Email: info@senhuang.org',
    size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))

# 修正欄寬（必須在所有 merge 完成後呼叫）
fix_table_col_widths(table, [8.23, 4.12, 3.96])

doc.save(str(OUT))
print(f"✓ 模板已生成：{OUT}")
