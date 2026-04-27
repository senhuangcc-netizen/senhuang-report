"""
東方森煌報告生成服務 — Railway 雲端版
POST /generate  →  下載照片 + XRF PDF → 生成 .docx → 上傳 Vercel Blob → 回傳下載連結
"""
from __future__ import annotations

import os
import shutil
import tempfile
import base64
import mimetypes
from datetime import datetime
from pathlib import Path

import io
import requests
import fitz  # PyMuPDF
from PIL import Image
from docx import Document
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import JSONResponse

app = FastAPI()

SECRET        = os.environ.get("WEBHOOK_SECRET", "")
ANTHROPIC_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
CLAUDE_MODEL  = os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6")

BLOB_TOKEN = os.environ.get("BLOB_READ_WRITE_TOKEN", "")

BASE_DOCX = Path(__file__).parent / "templates" / "base.docx"

PDF_KEYWORDS     = ["定量结果", "定量結果"]
PHOTO_EXTS       = {".jpg", ".jpeg", ".png", ".heic", ".webp"}


# ── 下載工具 ───────────────────────────────────────────────

def download(url: str, dest: Path) -> Path | None:
    if not url:
        return None
    try:
        r = requests.get(url, timeout=60)
        r.raise_for_status()
        dest.write_bytes(r.content)
        return dest
    except Exception as e:
        print(f"[下載失敗] {url}: {e}")
        return None


# ── PDF 截圖（原邏輯移植） ─────────────────────────────────

def screenshot_pdf_keyword(pdf_path: Path, out_path: Path) -> bool:
    doc = fitz.open(str(pdf_path))
    try:
        for page in doc:
            title = None
            for kw in PDF_KEYWORDS:
                hits = page.search_for(kw)
                if hits:
                    title = hits[0]
                    break
            if not title:
                continue
            elem_hits = [r for r in page.search_for("元素") if r.y0 >= title.y1 - 1]
            if not elem_hits:
                continue
            header = elem_hits[0]
            sigma_hits = page.search_for("3-sigma")
            if not sigma_hits:
                continue
            sig = min(sigma_hits, key=lambda r: abs(r.y0 - header.y0))
            # 右邊界：3-sigma 欄右側與「処理」欄左側的中線
            shori_hits = page.search_for("処理")
            if shori_hits:
                shori = min(shori_hits, key=lambda r: abs(r.y0 - header.y0))
                right_x = (sig.x1 + shori.x0) / 2
            else:
                right_x = sig.x1 + 2
            # 底部：Pb 最後一筆，否則用「定量分析-FP」為上緣，再否則用預設高度
            pb_hits = [r for r in page.search_for("Pb")
                       if r.x0 < header.x1 + 20 and r.y0 > header.y1]
            if pb_hits:
                bottom_y = max(r.y1 for r in pb_hits)
            else:
                fp_hits = page.search_for("定量分析-FP")
                if fp_hits:
                    bottom_y = min(r.y0 for r in fp_hits) - 2
                else:
                    bottom_y = header.y0 + 480
            clip = fitz.Rect(header.x0, header.y0, right_x, bottom_y + 1)

            # 結果欄位為 ND 的列，找出像素座標後用 Pillow 徹底移除
            # 使用 get_text("words") 做精確大小寫比對，避免誤中元素符號 Nd（釹）
            words = page.get_text("words")  # (x0,y0,x1,y1,word,block,line,idx)
            nd_hits = [fitz.Rect(w[0], w[1], w[2], w[3]) for w in words
                       if w[4].strip() == "ND"
                       and w[0] >= header.x0 - 5 and w[2] <= sig.x0 + 5
                       and w[1] >= header.y1 - 1 and w[3] <= bottom_y + 1]

            pix = page.get_pixmap(clip=clip, dpi=150)
            if pix.n != 3:
                pix = fitz.Pixmap(fitz.csRGB, pix)

            if nd_hits:
                scale = 150 / 72
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                img_h = img.height

                # 計算要刪除的像素行範圍（合併重疊區間）
                nd_rows: list[list[int]] = []
                for nd in nd_hits:
                    py0 = max(0, int((nd.y0 - clip.y0) * scale) - 1)
                    py1 = min(img_h, int((nd.y1 - clip.y0) * scale) + 2)
                    if nd_rows and py0 <= nd_rows[-1][1]:
                        nd_rows[-1][1] = max(nd_rows[-1][1], py1)
                    else:
                        nd_rows.append([py0, py1])

                # 保留的像素行段落
                keep: list[tuple[int, int]] = []
                prev = 0
                for y0, y1 in nd_rows:
                    if prev < y0:
                        keep.append((prev, y0))
                    prev = y1
                if prev < img_h:
                    keep.append((prev, img_h))

                # 拼接保留段落
                strips = [img.crop((0, y0, img.width, y1)) for y0, y1 in keep]
                total_h = sum(s.height for s in strips)
                new_img = Image.new("RGB", (img.width, total_h), (255, 255, 255))
                offset = 0
                for s in strips:
                    new_img.paste(s, (0, offset))
                    offset += s.height
                new_img.save(str(out_path), "PNG")
            else:
                pix.save(str(out_path))
            return True
    finally:
        doc.close()
    return False


# ── Claude 重寫鑑定說明 ────────────────────────────────────

def rewrite_description(category: str, ref_desc: str, xrf_png: Path) -> str:
    from anthropic import Anthropic
    client = Anthropic(api_key=ANTHROPIC_KEY)
    content: list = [{
        "type": "text",
        "text": (
            f"你是東方森煌古物鑑定所的鑑定專家。以下為同形制「{category}」"
            f"的參考鑑定說明，以及本件送檢品的 XRF 元素分析數據圖。\n\n"
            f"請依元素組成重寫一段約 180 字的鑑定說明，語氣專業、"
            f"以科學證據為基礎、繁體中文。\n\n"
            f"【參考鑑定說明】\n{ref_desc or '（無參考，請依圖直接撰寫）'}\n"
        ),
    }]
    if xrf_png.exists():
        content.append({
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": "image/png",
                "data": base64.standard_b64encode(xrf_png.read_bytes()).decode(),
            },
        })
    resp = client.messages.create(
        model=CLAUDE_MODEL, max_tokens=1024,
        messages=[{"role": "user", "content": content}],
    )
    return resp.content[0].text.strip()


# ── docx 圖片替換 ──────────────────────────────────────────

def replace_inline_image(doc: Document, index: int, new_path: Path) -> bool:
    shapes = doc.inline_shapes
    if index >= len(shapes):
        return False
    shape = shapes[index]
    rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
    part = doc.part.related_parts[rId]
    part._blob = new_path.read_bytes()
    mt = mimetypes.guess_type(str(new_path))[0] or "image/jpeg"
    try:
        part._content_type = mt
    except Exception:
        pass
    return True


def replace_run_text(para, new_text: str):
    runs = para.runs
    last = next((i for i in range(len(runs) - 1, -1, -1) if runs[i].text), -1)
    if last < 0:
        para.add_run(new_text)
    else:
        runs[last].text = new_text


def replace_value_after_colon_inline(para, new_text: str):
    runs = para.runs
    target = next((i for i in range(len(runs) - 1, -1, -1) if "：" in runs[i].text), -1)
    if target < 0:
        replace_run_text(para, new_text)
        return
    r = runs[target]
    prefix = r.text.split("：")[0] + "："
    r.text = prefix + new_text
    for j in range(target + 1, len(runs)):
        runs[j].text = ""


def replace_after_colon(para, new_text: str):
    runs = para.runs
    colon_idx = next((i for i, r in enumerate(runs) if "：" in r.text), -1)
    if colon_idx < 0:
        replace_run_text(para, new_text)
        return
    for j in range(colon_idx + 1, len(runs)):
        if runs[j].text and runs[j].text != "\n":
            runs[j].text = new_text
            return
    para.add_run(new_text)


def set_value_cell(cell, value: str):
    paras = cell.paragraphs
    if len(paras) >= 2:
        replace_run_text(paras[1], value)
    elif paras:
        replace_run_text(paras[0], value)


def generate_from_base(base_docx: Path, out_path: Path, fields: dict, images: dict):
    shutil.copy2(str(base_docx), str(out_path))
    doc = Document(str(out_path))

    for para in doc.paragraphs:
        t = para.text
        if "送驗編號" in t and "：" in t:
            replace_after_colon(para, fields.get("編號", ""))
        elif "顧客推估" in t and "：" in t:
            replace_after_colon(para, fields.get("形制", ""))

    if doc.tables:
        t0 = doc.tables[0]
        if t0.rows and len(t0.rows[0].cells) >= 2:
            replace_value_after_colon_inline(t0.rows[0].cells[0].paragraphs[0], fields.get("送檢日期", ""))
            replace_value_after_colon_inline(t0.rows[0].cells[1].paragraphs[0], fields.get("報告日期", ""))

    if len(doc.tables) >= 2:
        t1 = doc.tables[1]
        rows = t1.rows
        if len(rows) > 2:
            set_value_cell(rows[2].cells[0], fields.get("尺寸", "").replace(" mm", "").strip())
            set_value_cell(rows[2].cells[2], str(fields.get("重量", "")))
        if len(rows) > 3:
            set_value_cell(rows[3].cells[0], fields.get("材質", ""))
            set_value_cell(rows[3].cells[1], fields.get("形制", ""))
        if len(rows) > 4:
            set_value_cell(rows[4].cells[0], fields.get("說明", ""))
        if len(rows) > 5:
            set_value_cell(rows[5].cells[0], fields.get("鑑定結果", ""))
            set_value_cell(rows[5].cells[2], fields.get("備註", ""))

    order = [images.get("front"), images.get("xrf"), images.get("micro1"), images.get("micro2")]
    for i, img in enumerate(order):
        if img and Path(img).exists():
            replace_inline_image(doc, i, Path(img))

    doc.save(str(out_path))


# ── Vercel Blob 上傳 ──────────────────────────────────────

def upload_blob(file_path: Path, content_type: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document") -> str:
    date_prefix = datetime.now().strftime("%Y%m%d")
    pathname = f"reports/{date_prefix}/{file_path.name}"
    res = requests.put(
        f"https://blob.vercel-storage.com/{pathname}",
        data=file_path.read_bytes(),
        headers={
            "Authorization": f"Bearer {BLOB_TOKEN}",
            "Content-Type": content_type,
            "x-content-type": content_type,
        },
        timeout=60,
    )
    res.raise_for_status()
    return res.json()["url"]


# ── API ────────────────────────────────────────────────────

@app.post("/crop-xrf")
async def crop_xrf(request: Request):
    if request.headers.get("x-webhook-secret") != SECRET:
        raise HTTPException(status_code=401, detail="Unauthorized")
    body = await request.json()
    pdf_url = body.get("pdf_url", "")
    pdf_b64 = body.get("pdf_b64", "")
    if not pdf_url and not pdf_b64:
        raise HTTPException(status_code=400, detail="缺少 pdf_url 或 pdf_b64")
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)
        pdf_path = tmp / "xrf.pdf"
        if pdf_b64:
            pdf_path.write_bytes(base64.b64decode(pdf_b64))
        else:
            if not download(pdf_url, pdf_path):
                raise HTTPException(status_code=400, detail="PDF 下載失敗")
        xrf_png = tmp / "xrf_chart.png"
        ok = screenshot_pdf_keyword(pdf_path, xrf_png)
        if not ok:
            raise HTTPException(status_code=422, detail="找不到定量結果表格，請確認 PDF 含有「定量结果」欄位")
        img_b64 = base64.standard_b64encode(xrf_png.read_bytes()).decode()
    return JSONResponse({"image_b64": img_b64})


@app.get("/health")
def health():
    return {"status": "ok", "base_docx": BASE_DOCX.exists()}


@app.post("/generate")
async def generate(request: Request):
    if request.headers.get("x-webhook-secret") != SECRET:
        raise HTTPException(status_code=401, detail="Unauthorized")

    body = await request.json()
    folder_name  = body.get("folder_name", "report")
    item_code    = body.get("item_code", "")
    category_data = body.get("category_data", {})

    if not BASE_DOCX.exists():
        raise HTTPException(status_code=500, detail="base.docx 尚未放入 templates/ 資料夾")

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)

        front  = download(body.get("photo_front_url"),  tmp / "front.jpg")
        micro1 = download(body.get("photo_micro1_url"), tmp / "micro1.jpg")
        micro2 = download(body.get("photo_micro2_url"), tmp / "micro2.jpg")
        xrf_pdf = download(body.get("xrf_pdf_url"),     tmp / "xrf.pdf")

        xrf_png = tmp / "xrf.png"
        if xrf_pdf:
            ok = screenshot_pdf_keyword(xrf_pdf, xrf_png)
            if not ok:
                print("[警告] XRF PDF 找不到定量結果關鍵字")

        # 從 category_data 提取形制 / 材質
        形制_parts = [str(v) for k, v in category_data.items()
                     if k.startswith("形制_") and v]
        形制 = " ".join(形制_parts)
        材質_raw = category_data.get("材質", [])
        材質 = "、".join(材質_raw) if isinstance(材質_raw, list) else str(材質_raw)

        說明 = ""
        if ANTHROPIC_KEY:
            try:
                說明 = rewrite_description(形制, "", xrf_png)
            except Exception as e:
                print(f"[Claude 失敗] {e}")

        fields = {
            "編號":    item_code,
            "送檢日期": body.get("submission_date", ""),
            "報告日期": body.get("report_date", datetime.now().strftime("%d/%b/%Y").upper()),
            "尺寸":    body.get("size", ""),
            "重量":    body.get("weight", ""),
            "材質":    材質,
            "形制":    形制,
            "說明":    說明,
            "鑑定結果": body.get("appraisal_result", ""),
            "備註":    body.get("note", ""),
        }

        out_name = f"{形制.split()[0] if 形制 else folder_name}_{item_code}.docx"
        out_path = tmp / out_name
        generate_from_base(BASE_DOCX, out_path, fields, {
            "front":  str(front)  if front  else None,
            "xrf":    str(xrf_png) if xrf_png.exists() else None,
            "micro1": str(micro1) if micro1 else None,
            "micro2": str(micro2) if micro2 else None,
        })

        report_url = upload_blob(out_path)

    return JSONResponse({"success": True, "report_url": report_url})
